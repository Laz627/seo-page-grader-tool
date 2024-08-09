import streamlit as st
import pandas as pd
from openai import OpenAI
import os
import io
from docx import Document
import markdown
from bs4 import BeautifulSoup

def get_openai_api_key():
    if 'openai_api_key' not in st.session_state:
        st.session_state.openai_api_key = ''
    
    api_key = st.sidebar.text_input(
        "Enter your OpenAI API key",
        value=st.session_state.openai_api_key,
        type="password",
        key="openai_api_key_input"
    )
    
    if api_key:
        st.session_state.openai_api_key = api_key
        os.environ["OPENAI_API_KEY"] = api_key
        return api_key
    return None

# Set up the OpenAI API key prompt
st.sidebar.title("Setup")
api_key = get_openai_api_key()

# Initialize OpenAI client only if API key is provided
client = OpenAI() if api_key else None

# Define the SEO factors with criteria, explanations, and weights
seo_factors = {
    "On-Page": {
        "H1 Tag": {
            "criteria": [
                (
                    "Is included on-page at top of heading hierarchy",
                    10,
                    "The H1 tag is the main heading of the page and should be placed at the top of the content. It's crucial for SEO as it helps search engines understand the main topic of the page. To check this, view the page source and look for the <h1> tag, ensuring it's the first heading in the HTML structure.",
                ),
                (
                    "Contains proper length",
                    8,
                    "The H1 tag should be between 20-70 characters long. This length ensures it's descriptive enough for search engines while remaining concise for users. You can check the length by copying the H1 text and using a character counter tool.",
                ),
                (
                    "Contains primary keyword",
                    9,
                    "Including the primary keyword in the H1 tag helps search engines understand the main topic of the page. It should be placed naturally within the H1 text. To verify, check if your target keyword appears in the H1 tag.",
                ),
                (
                    "There is only a single H1 tag on-page",
                    8,
                    "Having only one H1 tag per page is a best practice for SEO. Multiple H1 tags can confuse search engines about the main topic of the page. To check, use the browser's 'Inspect' tool and search for 'h1' in the HTML.",
                ),
            ]
        },
        "Meta Title": {
            "criteria": [
                (
                    "Contains proper length",
                    9,
                    "The meta title should be 50-60 characters long. This ensures it's fully displayed in search results without being cut off. You can check the length using SEO tools like Moz's Title Tag Preview tool.",
                ),
                (
                    "Contains primary keyword",
                    10,
                    "Including the primary keyword in the meta title helps search engines understand what the page is about. Place the keyword naturally near the beginning of the title. You can view the meta title in the page source or using SEO browser extensions.",
                ),
                (
                    "There is only a single meta title on-page",
                    8,
                    "Each page should have only one meta title tag. Multiple title tags can confuse search engines. Check the page source to ensure there's only one <title> tag in the <head> section.",
                ),
            ]
        },
        "Meta Description": {
            "criteria": [
                (
                    "Contains proper length",
                    5,
                    "The ideal meta description length is between 150-160 characters. This length allows for a comprehensive summary without being cut off in search results. Use SEO tools or character counters to check the length.",
                ),
                (
                    "Contains primary keyword",
                    6,
                    "Including the primary keyword in the meta description helps improve click-through rates from search results. Place the keyword naturally within the description. You can view the meta description in the page source or using SEO browser extensions.",
                ),
                (
                    "There is only a single meta description on-page",
                    4,
                    "Each page should have only one meta description. Multiple descriptions can lead to inconsistent display in search results. Check the page source to ensure there's only one meta description tag.",
                ),
                (
                    "Adequately describes the purpose of the page as a CTA",
                    5,
                    "The meta description should act as a call-to-action, enticing users to click through to your page. It should clearly state what the page offers and why users should visit. Evaluate if your description compels action and accurately represents the page content.",
                ),
            ]
        },
        "Proper Heading Hierarchy": {
            "criteria": [
                (
                    "Only a single H1; H2s follow H1 tag; H3s follow H2s, etc…",
                    7,
                    "Proper heading hierarchy helps search engines understand the structure and importance of your content. Use H1 for the main title, H2 for main sections, H3 for subsections, and so on. To check, use the browser's 'Inspect' tool to view the HTML structure and ensure headings are nested correctly.",
                )
            ]
        },
        "Image Alt Text": {
            "criteria": [
                (
                    "Images include alt text with target keyword",
                    3,
                    "Alt text describes images for search engines and visually impaired users. Include your target keyword naturally in the alt text when relevant. To check, inspect the image HTML and look for the 'alt' attribute.",
                ),
                (
                    "Alt text properly describes imagery in a meaningful way",
                    2,
                    "Alt text should accurately describe the image content, not just repeat keywords. This improves accessibility and helps search engines understand the context. Review each image's alt text to ensure it provides a clear, concise description.",
                ),
            ]
        },
        "Schema Markup": {
            "criteria": [
                (
                    "Schema is included on-page as JSON-LD",
                    9,
                    "Schema markup helps search engines understand your content better, potentially leading to rich snippets in search results. Use JSON-LD format for easy implementation. To check, view the page source and look for a script tag containing JSON-LD data.",
                ),
                (
                    "No errors or warnings with schema markup",
                    8,
                    "Errors in schema markup can prevent search engines from using it effectively. Use Google's Structured Data Testing Tool to check for errors and warnings in your schema implementation.",
                ),
                (
                    "Schema markup matches page intent",
                    9,
                    "The schema type should accurately represent the page content (e.g., Article, Product, FAQ). Mismatched schema can confuse search engines. Review your schema type and ensure it aligns with the page's primary purpose.",
                ),
            ]
        },
        "Internal Linking": {
            "criteria": [
                (
                    "Other pages properly point to this one with target keyword included in anchor text",
                    7,
                    "Internal links with relevant anchor text help search engines understand the context and importance of the linked page. Check other pages on your site and ensure they link to this page using descriptive, keyword-rich anchor text when appropriate.",
                ),
                (
                    "This page logically drives users to the next anticipated step in the user journey",
                    6,
                    "Internal links should guide users through a logical flow on your website. Analyze the user journey and ensure this page links to the next logical step or related content that users might be interested in.",
                ),
                (
                    "This page doesn't send users to a dead end experience / poor off-ramp",
                    5,
                    "Every page should provide clear next steps for users. Ensure this page has relevant internal links, calls-to-action, or related content suggestions to keep users engaged and on your site.",
                ),
            ]
        },
        "User Engagement Metrics": {
            "criteria": [
                (
                    "Bounce rate meets or exceeds baseline",
                    8,
                    "A low bounce rate indicates that users find the content relevant. Compare the page's bounce rate to your site average or industry benchmarks. You can find this data in Google Analytics or similar analytics tools.",
                ),
                (
                    "Time spent on-page meets or exceeds baseline",
                    9,
                    "Longer time on page suggests engaging content. Compare the average time on this page to your site average or industry benchmarks. This data is available in most analytics platforms.",
                ),
                (
                    "CTR / average KW position meets or exceeds baseline",
                    8,
                    "High CTR and good keyword positions indicate effective optimization. Check Google Search Console for these metrics and compare them to your site averages or industry benchmarks.",
                ),
                (
                    "Visits meet or exceed baseline",
                    7,
                    "Higher visit numbers suggest the page is attracting significant traffic. Compare the page's visit count to your site average or set goals. This data is available in your analytics platform.",
                ),
                (
                    "Conversions / abandons meet or exceed baseline",
                    8,
                    "Good conversion rates indicate the page effectively meets user intent. Set up conversion tracking in your analytics tool and compare this page's performance to your site average or industry benchmarks.",
                ),
                (
                    "Scroll depth meets or exceeds baseline",
                    6,
                    "Greater scroll depth indicates users are consuming more content. Use scroll depth tracking in your analytics tool to measure how far users scroll on the page, and compare it to your site average.",
                ),
            ]
        },
        "Primary Topic/Keyword Targeting": {
            "criteria": [
                (
                    "Keyword is included above the fold in content",
                    10,
                    "Including the primary keyword early in the content helps search engines quickly understand the page topic. Check if your keyword appears in the first paragraph or section of your content, visible without scrolling.",
                ),
                (
                    "Relevant secondary keywords are included within subheads / body copy of page",
                    7,
                    "Secondary keywords help establish topical relevance and can help you rank for related terms. Use keyword research tools to identify relevant secondary keywords, then naturally incorporate them into your subheadings and body text.",
                ),
                (
                    "Page matches expected keyword intent",
                    9,
                    "Content should align with the user's search intent (informational, navigational, or transactional). Analyze the top-ranking pages for your target keyword to understand the intent, and ensure your content matches it.",
                ),
            ]
        },
        "URL Slug": {
            "criteria": [
                (
                    "Short length",
                    3,
                    "Shorter URLs are easier to read, share, and remember. Aim for 3-5 words in your URL slug. You can check and edit this in your CMS or website platform.",
                ),
                (
                    "Omission of stop words",
                    2,
                    "Removing unnecessary words (like 'the', 'a', 'an') makes URLs cleaner and more focused. Review your URL and remove any stop words that don't add value.",
                ),
                (
                    "Aligns with informational architecture of domain",
                    4,
                    "URLs should reflect your site's structure. Ensure the URL fits logically within your site's hierarchy. This can usually be set in your CMS or website platform.",
                ),
                (
                    "Lowercase only",
                    2,
                    "Using only lowercase letters in URLs helps avoid duplicate content issues. Most CMS automatically enforce this, but check to ensure all letters in your URL are lowercase.",
                ),
                (
                    "Hyphens only",
                    2,
                    "Use hyphens (-) instead of underscores (_) or spaces to separate words in URLs. This improves readability for both users and search engines. Check your URL structure and replace any underscores or spaces with hyphens.",
                ),
                (
                    "Non-parameterized (optional)",
                    1,
                    "Clean URLs without parameters are preferred. If possible, avoid query strings (e.g., '?id=123') in your URLs. This may require adjustments to your website's configuration.",
                ),
                (
                    "ASCII characters only",
                    1,
                    "Stick to standard ASCII characters in URLs for best compatibility. Avoid special characters or non-English letters. This is usually handled automatically by most CMS.",
                ),
                (
                    "Depth of 5 or less from the homepage",
                    2,
                    "Keeping URLs close to the homepage in the site structure can boost their perceived importance. Try to keep your URL structure no more than 5 levels deep. This may require reorganizing your site structure.",
                ),
            ]
        },
        "Quality of Content": {
            "criteria": [
                (
                    "Accuracy",
                    10,
                    "Ensure all information is factual and up-to-date. Regularly review and update your content. Use credible sources and link to them when appropriate.",
                ),
                (
                    "Originality",
                    9,
                    "Create unique content that adds value. Avoid duplicating content from other sources. Use plagiarism checkers to ensure your content is original.",
                ),
                (
                    "Tone of voice matches brand standards",
                    8,
                    "Maintain a consistent brand voice across all content. Develop and follow brand guidelines for tone and style. Regularly review content to ensure it aligns with your brand voice.",
                ),
                (
                    "Topic completeness",
                    10,
                    "Cover the topic comprehensively. Research competing content and ensure you're addressing all relevant aspects of the topic. Consider using topic clustering techniques to cover subjects thoroughly.",
                ),
                (
                    "Readability",
                    8,
                    "Content should be easy to read and understand. Use tools like the Flesch-Kincaid readability test to assess and improve your content's readability. Break up text with subheadings, short paragraphs, and bullet points.",
                ),
                (
                    "Formatting (paragraph breaks, logical subheading structure)",
                    7,
                    "Well-formatted content improves user experience and readability. Use short paragraphs, clear subheadings, and appropriate use of bold and italic text. Ensure your content has a logical structure that's easy to follow.",
                ),
                (
                    "Content freshness / regularly updated",
                    9,
                    "Keep your content current and relevant. Regularly update your content with new information, examples, or data. Add a 'last updated' date to your pages to show freshness.",
                ),
                (
                    "Page matches expected user intent",
                    10,
                    "Ensure your content aligns with what users are looking for when they search for your target keywords. Analyze top-ranking pages and user feedback to understand and match user intent.",
                ),
                (
                    "Other pages don't cannibalize this one for content",
                    8,
                    "Avoid having multiple pages competing for the same keywords. Conduct a content audit to identify and resolve any cannibalization issues. Consider consolidating similar content or using canonical tags where appropriate.",
                ),
            ]
        },
    },
    "Off-Page": {
        "Page Authority vs Top 10": {
            "criteria": [
                (
                    "Page authority is greater than average of top 10 results",
                    7,
                    "Page Authority (PA) is a metric that predicts how well a page will rank. Use tools like Moz to check your PA and compare it to the average PA of the top 10 results for your target keyword. If your PA is lower, focus on building high-quality backlinks to improve it.",
                )
            ]
        },
        "Page Authority vs Top 3": {
            "criteria": [
                (
                    "Page authority is greater than average of top 3 results",
                    9,
                    "Comparing your PA to the top 3 results gives you a benchmark for highly competitive positions. Use SEO tools to check the PA of the top 3 results and aim to match or exceed their average. This often requires a strong backlink profile and excellent on-page optimization.",
                )
            ]
        },
        "Backlinks from Relevant Domains": {
            "criteria": [
                (
                    "Backlinks are from topically relevant domains",
                    7,
                    "Links from sites in your industry or niche carry more weight. Use backlink analysis tools to check the relevance of your linking domains. Focus on acquiring links from sites that are topically related to your content.",
                )
            ]
        },
        "Backlink Placement": {
            "criteria": [
                (
                    "Backlinks are placed higher up on sourced pages / are likely to be clicked",
                    3,
                    "Links placed prominently on a page (e.g., in the main content area rather than the footer) are more valuable. Analyze your backlinks to see where they appear on the linking pages. Aim for contextual links within the main content of high-quality pages.",
                )
            ]
        },
        "Backlink Anchor Text": {
            "criteria": [
                (
                    "Backlinks contain topically relevant anchor text",
                    6,
                    "Anchor text helps search engines understand what the linked page is about. Analyze your backlinks' anchor text using SEO tools. Aim for a natural mix of branded, keyword-rich, and generic anchor texts, avoiding over-optimization.",
                )
            ]
        },
        "Backlink Traffic": {
            "criteria": [
                (
                    "Backlinks are placed on pages that actually drive visits",
                    7,
                    "Links from pages with high traffic can increase your visibility and drive referral traffic. Use tools like Ahrefs to estimate the traffic of pages linking to you. Focus on acquiring links from popular, high-traffic pages in your niche.",
                )
            ]
        },
    },
    "Technical": {
        "Canonical Tag": {
            "criteria": [
                (
                    "Canonical tag contains self-reference",
                    6,
                    "The canonical tag tells search engines which version of a page is the preferred one to index. Check your page's HTML for the canonical tag in the <head> section. Ensure it points to the current URL to avoid indexing issues.",
                )
            ]
        },
        "Hreflang Tag": {
            "criteria": [
                (
                    "Hreflang tag (optional) is correct, targets the right locations, and references other translated page equivalents",
                    6,
                    "Hreflang tags help search engines serve the correct language version of your page. If your site has multiple language versions, check the <head> section for correct hreflang implementation. Ensure each language version is properly referenced.",
                )
            ]
        },
        "Indexability": {
            "criteria": [
                (
                    "Page is indexable by search engines / isn't blocked by meta tags or robots.txt",
                    10,
                    "For a page to appear in search results, it must be indexable. Check your robots.txt file and the page's HTML for any 'noindex' directives. Use Google Search Console's URL Inspection tool to verify if the page is indexable.",
                )
            ]
        },
        "Sitemap Inclusion": {
            "criteria": [
                (
                    "Page is included in sitemap.xml file",
                    2,
                    "Sitemaps help search engines discover and understand your site structure. Check your sitemap.xml file to ensure the page is listed. You can usually find your sitemap at yourdomain.com/sitemap.xml or check your CMS settings.",
                )
            ]
        },
        "Page Orphan Status": {
            "criteria": [
                (
                    "Page isn't orphaned",
                    2,
                    "Orphan pages are not linked to from any other page on your site, making them hard to find. Use a site crawling tool to identify orphan pages. Ensure all important pages are linked to from at least one other page on your site.",
                )
            ]
        },
        "Renderability": {
            "criteria": [
                (
                    "Page elements are renderable by search engines",
                    6,
                    "Search engines should be able to render all important content on your page. Use Google Search Console's URL Inspection tool to see how Googlebot renders your page. Ensure all critical content is visible and not reliant on JavaScript that might not be executed by search engine crawlers.",
                )
            ]
        },
        "Web Core Vitals": {
            "criteria": [
                (
                    "Page passes Web Core Vitals metrics / exceeds industry average",
                    3,
                    "Web Core Vitals are a set of metrics that measure user experience in terms of loading performance, interactivity, and visual stability. Use Google PageSpeed Insights or the Core Web Vitals report in Google Search Console to check your performance. Aim to have all Core Web Vitals in the 'good' range.",
                )
            ]
        },
        "Mobile Friendliness": {
            "criteria": [
                (
                    "Page is mobile-friendly and responsive",
                    8,
                    "With mobile-first indexing, having a mobile-friendly site is crucial. Use Google's Mobile-Friendly Test tool to check your page. Ensure your site uses responsive design and provides a good user experience on all device sizes.",
                )
            ]
        },
        "HTTPS": {
            "criteria": [
                (
                    "Page is served over HTTPS",
                    5,
                    "HTTPS is a ranking factor and provides security for your users. Check if your URL starts with 'https://'. If not, obtain an SSL certificate and implement HTTPS across your entire site.",
                )
            ]
        },
    },
}

# Calculate bucket weights
bucket_weights = {
    "On-Page": 0.55,
    "Off-Page": 0.30,
    "Technical": 0.15
}

import streamlit as st
import pandas as pd
from openai import OpenAI
import os
from docx import Document
import io

def get_user_input(factor, criteria):
    responses = {}
    st.subheader(factor)
    
    for i, (criterion, weight, help_text) in enumerate(criteria):
        col1, col2 = st.columns([4, 1])
        with col1:
            options = ["Yes", "No", "N/A"] if "optional" in criterion.lower() else ["Yes", "No"]
            responses[criterion] = {
                "response": st.radio(criterion, options, index=1, key=f"{factor}_{i}"),
                "weight": weight
            }
        with col2:
            st.markdown(f'''
                <div title="{help_text}" style="
                    cursor: help;
                    background-color: #e1e4e8;
                    color: #0366d6;
                    width: 20px;
                    height: 20px;
                    border-radius: 50%;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    font-weight: bold;
                    font-size: 14px;
                    margin: auto;
                ">?</div>
            ''', unsafe_allow_html=True)
    
    st.markdown("<hr style='margin-top: 20px; margin-bottom: 20px;'>", unsafe_allow_html=True)
    
    return responses

def calculate_score(inputs):
    score = 0
    max_score = 0
    for criterion, data in inputs.items():
        if data["response"] == "Yes":
            score += data["weight"]
        if data["response"] != "N/A":
            max_score += data["weight"]
    return score / max_score * 10 if max_score > 0 else 0

def estimate_ranking(overall_score):
    if overall_score >= 9.5:
        return "1-3"
    elif overall_score >= 9.0:
        return "4-6"
    elif overall_score >= 8.5:
        return "7-10"
    elif overall_score >= 8.0:
        return "11-15"
    elif overall_score >= 7.5:
        return "16-20"
    elif overall_score >= 7.0:
        return "21-30"
    elif overall_score >= 6.5:
        return "31-50"
    elif overall_score >= 6.0:
        return "51-100"
    else:
        return "100+"

def get_gpt4_recommendations(inputs):
    st.write("Generating recommendations with OpenAI's GPT-4...")
    prompt = f"""Based on the following SEO audit results, provide recommendations for improvement:

{inputs}

Please provide specific, actionable recommendations for each area that needs improvement. Use the following format:

### [Main Category]
**[Subcategory]**
- Action: [Specific recommendation]
- Action: [Another specific recommendation]

Repeat this structure for each category and subcategory that needs improvement."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an SEO expert providing recommendations based on an audit."},
                {"role": "user", "content": prompt}
            ]
        )
        st.write("Recommendations generated successfully!")
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error generating recommendations: {str(e)}")
        return "Unable to generate recommendations at this time."

def export_to_word(inputs, scores, recommendations, estimated_ranking):
    doc = Document()
    doc.add_heading('SEO Audit Results', 0)

    # Scores
    doc.add_heading('Scores', level=1)
    for bucket, score in scores.items():
        doc.add_paragraph(f"{bucket}: {score:.2f}/10")
    
    # Estimated Ranking
    doc.add_heading('Estimated Ranking', level=1)
    doc.add_paragraph(f"Based on the overall score of {scores['Overall']:.2f}/10, the page might rank in positions: {estimated_ranking}")

    # Recommendations
    doc.add_heading('Recommendations', level=1)
    
    # Split recommendations into lines
    rec_lines = recommendations.split('\n')
    
    current_heading = None
    for line in rec_lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('###'):
            # Main heading
            doc.add_heading(line.strip('# '), level=2)
            current_heading = None
        elif line.startswith('**') and line.endswith('**'):
            # Subheading
            doc.add_heading(line.strip('*'), level=3)
            current_heading = line.strip('*')
        elif line.startswith('-') or line.startswith('*'):
            # Bullet point
            doc.add_paragraph(line.strip('- *'), style='List Bullet')
        elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
            # Numbered list
            doc.add_paragraph(line, style='List Number')
        else:
            # Regular paragraph
            if current_heading:
                doc.add_paragraph(f"{current_heading}: {line}")
            else:
                doc.add_paragraph(line)

    # Selected Criteria
    doc.add_heading('Selected Criteria', level=1)
    for bucket, factors in seo_factors.items():
        doc.add_heading(f"{bucket} Factors", level=2)
        for factor, criteria in factors.items():
            doc.add_heading(factor, level=3)
            for criterion, data in inputs[bucket][factor].items():
                if data['response'] != "N/A":
                    doc.add_paragraph(f"{criterion}: {data['response']}")

    # Save the document to a BytesIO object
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

def main():
    st.title("SEO Page Ranking Calculator")

    st.markdown("""
    Created by Brandon Lazovic

    This tool helps you evaluate the SEO potential of a web page by assessing various on-page, off-page, and technical factors. 
    To use the tool:
    1. Enter your OpenAI API key in the sidebar (for personalized recommendations).
    2. Go through each factor and select 'Yes' or 'No' based on whether your page meets the criteria.
    3. Click 'Calculate Score' to see your results and recommendations.
    4. Download a detailed report of your audit.
    
    The tool provides an overall score and individual scores for on-page, off-page, and technical factors.
    """)

    if not api_key:
        st.warning("Please enter your OpenAI API key in the sidebar to enable recommendations.")

    inputs = {}
    for bucket, factors in seo_factors.items():
        bucket_inputs = {}
        for factor, data in factors.items():
            bucket_inputs[factor] = get_user_input(factor, data["criteria"])
        inputs[bucket] = bucket_inputs

    if st.button("Calculate Score"):
        with st.spinner("Calculating scores..."):
            progress_bar = st.progress(0)
            scores = {}
            for i, (bucket, factors) in enumerate(seo_factors.items()):
                bucket_score = sum(calculate_score(inputs[bucket][factor]) for factor in factors)
                scores[bucket] = bucket_score / len(factors)
                progress_bar.progress((i + 1) / len(seo_factors))

            overall_score = sum(score * bucket_weights[bucket] for bucket, score in scores.items())
            scores["Overall"] = overall_score
            progress_bar.progress(100)

        st.subheader("SEO Scores")
        st.markdown("<div style='background-color: #e6f3ff; padding: 10px; border-radius: 5px;'>", unsafe_allow_html=True)
        for bucket, score in scores.items():
            st.write(f"{bucket}: {score:.2f}/10")
        st.markdown("</div>", unsafe_allow_html=True)

        estimated_ranking = estimate_ranking(scores["Overall"])
        st.subheader("Estimated Ranking")
        st.write(f"Based on your overall score of {scores['Overall']:.2f}/10, your page might rank in positions: {estimated_ranking}")
          
        st.markdown("""
      
        **Note:** This ranking estimate is a rough approximation based on your SEO score. Actual rankings can vary significantly due to factors such as:
        - Competition in your specific niche
        - Search intent alignment
        - Domain authority
        - Freshness of content
        - User engagement metrics
        - Regular algorithm updates
    
        Use this estimate as a general guide rather than a guaranteed outcome.
        """)

        if api_key:
            with st.spinner("Generating recommendations..."):
                recommendations = get_gpt4_recommendations(inputs)
            st.subheader("Recommendations")
            st.markdown("<div style='background-color: #e6ffe6; padding: 10px; border-radius: 5px;'>", unsafe_allow_html=True)
            st.write(recommendations)
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.warning("OpenAI API key not provided. Recommendations are not available.")

        with st.spinner("Preparing download..."):
            # Generate the Word document
            doc_bytes = export_to_word(inputs, scores, recommendations, estimated_ranking)

        # Provide a download button for the generated document
        st.download_button(
            label="Download SEO Audit Results",
            data=doc_bytes,
            file_name="seo_audit_results.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
